"""Generate the graduation thesis Word document with proper TJUT formatting."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12), bold=False):
    """Set font for a run with both Chinese and English font support."""
    run.font.size = size
    run.bold = bold
    run.font.name = en_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rPr.insert(0, rFonts)


def set_paragraph_spacing(paragraph, line_spacing=1.25, before=0, after=0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def add_heading_custom(doc, text, level=1):
    """Add a chapter/section heading with proper formatting.

    level 0: 章标题 (三号黑体, centered)
    level 1: 节标题 (小三号黑体, centered)
    level 2: 小节标题 (四号黑体, left-aligned)
    """
    sizes = {0: Pt(16), 1: Pt(15), 2: Pt(14)}  # 三号=16pt, 小三=15pt, 四号=14pt
    alignments = {0: WD_ALIGN_PARAGRAPH.CENTER, 1: WD_ALIGN_PARAGRAPH.CENTER, 2: WD_ALIGN_PARAGRAPH.LEFT}

    p = doc.add_paragraph()
    p.alignment = alignments[level]
    set_paragraph_spacing(p, line_spacing=1.5, before=12, after=6)
    run = p.add_run(text)
    set_run_font(run, cn_font='黑体', en_font='Times New Roman', size=sizes[level], bold=True)
    return p


def add_body_text(doc, text, first_line_indent=True):
    """Add body text (小四号宋体, 1.25倍行距)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, line_spacing=1.25, before=0, after=3)
    if first_line_indent:
        pf = p.paragraph_format
        pf.first_line_indent = Pt(24)  # ~2 Chinese characters
    run = p.add_run(text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12), bold=False)  # 小四号=12pt
    return p


def add_ref_text(doc, text):
    """Add reference item (小四号宋体)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, line_spacing=1.25, before=0, after=2)
    run = p.add_run(text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12), bold=False)
    return p


def set_page_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    """Set page margins in cm."""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def create_thesis():
    doc = Document()

    # Page setup
    set_page_margins(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # ======================================================================
    # 第一章 绪论
    # ======================================================================
    add_heading_custom(doc, '第一章 绪论', level=0)

    # ---------- 1.1 ----------
    add_heading_custom(doc, '1.1 研究背景与意义', level=1)

    add_body_text(doc, '目标检测作为计算机视觉领域的核心研究方向，旨在从图像或视频中定位并识别感兴趣的物体。近年来，随着深度学习技术的迅猛发展，基于卷积神经网络（Convolutional Neural Network, CNN）的目标检测算法在通用场景下取得了令人瞩目的成果。以YOLO（You Only Look Once）系列、R-CNN系列和DETR系列为代表的检测框架，在COCO（Common Objects in Context）、PASCAL VOC等大规模公开数据集上不断刷新性能纪录，检测精度和推理速度均达到了前所未有的水平。')

    add_body_text(doc, '然而，小目标检测始终是目标检测领域中最具挑战性的任务之一。根据COCO数据集的定义，小目标是指像素面积小于32×32的物体。在真实世界场景中，小目标普遍存在于各类应用中：无人机高空航拍图像中的车辆和行人、卫星遥感图像中的建筑物和舰船、交通监控视频中远处的行人、工业缺陷检测中的微小裂纹等。这些场景中，目标占图像总面积比例极小，包含的特征信息非常有限，导致常规检测算法难以有效捕捉和识别。')

    add_body_text(doc, '小目标检测面临的核心难题可归纳为以下几个方面：（1）特征表达能力不足——小目标在深层卷积网络中经过多次下采样后，特征图分辨率急剧下降，导致小目标的特征响应微弱甚至消失，网络难以提取到具有判别力的语义信息；（2）上下文信息利用不充分——小目标缺乏足够的局部上下文线索，而全局上下文关联又难以建立，使得定位和分类容易受到背景噪声干扰；（3）正负样本极端不平衡——图像中大部分区域为背景，小目标的正样本锚框数量远少于负样本，导致模型训练时偏向背景类别，抑制了对小目标的检测性能；（4）定位精度要求更高——小目标对边界框偏移极为敏感，微小的定位误差即可能导致IoU（Intersection over Union）大幅下降，触发漏检。')

    add_body_text(doc, '无人机航拍技术的快速发展为小目标检测研究提供了重要的应用场景和实验平台。近年来，无人机在环境监测、灾害救援、农业植保、交通管理、国防安全等领域展现出广阔的应用前景。无人机搭载高分辨率光学载荷，以灵活机动的视角获取大面积地物信息，能够有效弥补地面监测和卫星遥感的不足。然而，在实际应用中，为保障地面活动正常进行并扩大单次航程覆盖面积，无人机往往需要保持较高的飞行高度。这种作业特点导致航拍图像中的目标像素尺寸极小、空间分辨率有限，常规目标检测算法在此类数据上性能急剧退化。')

    add_body_text(doc, 'VisDrone2019数据集是由天津大学机器学习与数据挖掘实验室发布的大规模无人机航拍目标检测基准数据集，包含6471张训练图像、548张验证图像和1610张测试图像，涵盖行人、人、自行车、汽车、货车、卡车、三轮车、带篷三轮车、公交车和摩托车共10个类别。图像采集自中国多个城市，覆盖了不同天气条件、光照环境和场景类型，标注了超过260万个目标边界框。该数据集中的目标具有尺寸小、密集分布、尺度变化大、遮挡严重等特点，是评估小目标检测算法性能的理想平台。')

    add_body_text(doc, '深度学习目标检测算法在无人机平台上的部署同样面临严峻挑战。无人机搭载的嵌入式计算平台算力有限、存储空间紧凑、功耗约束严格，而大多数主流检测算法的模型参数量大、计算复杂度过高，难以在无人机端实现实时推理。因此，研究兼顾检测精度和模型效率的轻量化小目标检测算法，不仅具有重要的学术意义，也具有显著的工程应用价值。')

    add_body_text(doc, '综上所述，本文以无人机航拍场景下的小目标检测为研究对象，以VisDrone2019数据集为实验平台，以YOLOv11检测框架为技术基线，针对小目标特征表达不足、多尺度信息利用不充分、定位精度受限等关键问题，研究引入P2检测层扩充特征金字塔的多尺度感知能力，引入EMA（Efficient Multi-Scale Attention）注意力机制增强关键特征区域的响应，优化损失函数设计以改善小目标的回归精度。同时设计并实现基于PyQt5的可视化小目标检测系统，将改进算法转化为可交互的应用工具。研究成果对推动无人机航拍小目标检测技术的实际应用具有重要参考价值。')

    # ---------- 1.2 ----------
    add_heading_custom(doc, '1.2 国内外研究现状', level=1)

    # 1.2.1
    add_heading_custom(doc, '1.2.1 小目标检测研究现状', level=2)

    add_body_text(doc, '小目标检测作为目标检测领域的重要分支，长期以来受到国内外学者的广泛关注。传统的目标检测方法主要依赖手工设计的特征描述子（如SIFT、HOG等）与滑动窗口或选择性搜索相结合的策略，在小目标场景下性能严重受限。深度学习兴起后，基于卷积神经网络的特征自动学习范式逐步取代了手工特征时代，小目标检测性能取得了质的飞跃。')

    add_body_text(doc, '当前的小目标检测研究可归纳为以下几条技术路线：')

    add_body_text(doc, '（1）多尺度特征融合方法。通过构建特征金字塔结构，将高层的强语义特征与低层的高分辨率特征进行融合，提升小目标的特征表达能力。Lin等提出的特征金字塔网络（Feature Pyramid Network, FPN）是该方向的里程碑式工作，通过自上而下的路径和横向连接，将各层级的特征信息有效整合。Liu等在此基础上提出PANet（Path Aggregation Network），增加了自下而上的路径增强分支，进一步缩短了低层特征到高层特征的信息传递距离。Ghiasi等提出的NAS-FPN使用神经架构搜索自动设计特征金字塔拓扑结构。Tan等提出的BiFPN（Bi-directional FPN）引入加权特征融合机制和跨尺度连接，在EfficientDet中取得了突出的小目标检测性能。')

    add_body_text(doc, '（2）超分辨率增强方法。利用超分辨率重建技术提升小目标区域的特征质量，使其接近中大型目标的特征表达水平。Li等提出的Perceptual GAN通过生成对抗网络为小目标生成高分辨率特征表示，缩小大小目标之间的特征鸿沟。Noh等提出了一种基于特征级别超分辨率的方法，在特征空间内对小目标区域进行上采样和细化处理。Bai等提出的SOD-MTGAN利用多任务生成对抗网络同步完成检测和超分辨率重建。')

    add_body_text(doc, '（3）上下文信息利用方法。通过显式建模目标与周围场景的空间关系和语义关联，弥补小目标自身特征的不足。Hu等提出的关系网络（Relation Network）通过目标间的关系建模增强小目标的语义表征。Liu等提出的SINet利用场景上下文信息为小目标提供补充判别线索。Zhu等提出的Cascade R-CNN通过多级级联的IoU阈值递进优化，克服了小目标在高IoU阈值下正样本匮乏的问题。')

    add_body_text(doc, '（4）数据增强与训练策略。通过针对性的数据增强和训练技巧解决小目标样本不足和分布不均的问题。Kisantal等提出的小目标复制粘贴增强策略（Copy-Paste Augmentation），将小目标实例从训练图像中裁剪并粘贴到其他图像的合适位置，有效增加了小目标的样本数量和上下文多样性。Ghiasi等提出的简单复制粘贴策略在不增加计算开销的前提下显著提升小目标性能。Chen等提出的尺度感知训练策略（Scale-Aware Training）和多分辨率训练混合机制，使模型适应更广泛的尺度分布。')

    add_body_text(doc, '（5）无锚框检测方法。传统基于锚框的检测方法需要预设大量锚框参数，小目标的锚框匹配率往往很低。无锚框方法通过关键点热力图或中心点预测进行目标定位，消除了锚框设计带来的超参数敏感性。Zhou等提出的CenterNet将目标建模为中心点及其宽高属性，避免了复杂的锚框匹配过程。Law等提出的CornerNet通过检测目标左上角和右下角的角点对实现目标定位。Tian等提出的FCOS以逐像素预测的方式直接回归目标边界框，结合中心度（Center-ness）分支抑制低质量检测框。Kong等提出的FoveaBox通过预测目标中心区域的存在性和边界偏移实现无锚框检测。')

    add_body_text(doc, '尽管上述方法在通用数据集上取得了不同程度的提升，但在无人机航拍等极端小目标密度场景下，仍面临特征分辨率不足、多尺度退化严重、上下文建模复杂度高等挑战，需要进一步探索更有效的解决方案。')

    # 1.2.2
    add_heading_custom(doc, '1.2.2 YOLO系列算法演进', level=2)

    add_body_text(doc, 'YOLO（You Only Look Once）系列算法是目标检测领域最具影响力的单阶段检测框架，以其简洁的架构和优异的推理速度广泛应用于实际部署场景。自2016年Redmon等提出YOLOv1以来，该系列经历了超过十个版本的迭代演进，每一次更新都在网络结构、训练策略和检测范式上进行了重要改进。')

    add_body_text(doc, 'YOLOv1首次将目标检测统一为回归问题，通过将输入图像划分为S×S的网格单元，每个单元预测固定数量的边界框和类别概率，实现了端到端的一次性检测。该设计以45 FPS的推理速度大幅领先当时的双阶段检测方法，但在小目标和密集目标的检测精度上表现不佳。YOLOv2通过引入批量归一化（Batch Normalization）、高分辨率分类器预训练和锚框机制（Anchor Boxes），将检测精度显著提升。同时采用维度聚类（Dimension Clusters）通过K-means算法自动学习锚框先验尺寸，改善了尺度适应能力。')

    add_body_text(doc, 'YOLOv3将骨干网络从Darknet-19升级为Darknet-53，引入残差连接（Residual Connection）使网络能够训练得更深。特征融合方面，借鉴FPN的思想在三个不同尺度上进行预测，使网络能够同时检测大、中、小目标。分类头采用多个独立的二分类逻辑回归分类器替代Softmax，支持多标签分类场景。该版本在精度和速度之间取得了良好的平衡，标志着YOLO系列走向成熟。')

    add_body_text(doc, 'YOLOv4由Bochkovskiy等在2020年提出，系统性地总结了目标检测中的各种优化技巧。作者将改进分为两大类：仅增加训练成本而不增加推理时间的"免费午餐"（Bag of Freebies），包括Mosaic数据增强、CIoU损失函数、类标签平滑、余弦退火学习率调度等；稍微增加推理时间但显著提升精度的"特殊礼物"（Bag of Specials），包括Mish激活函数、跨阶段部分连接（Cross Stage Partial, CSP）、空间金字塔池化（Spatial Pyramid Pooling, SPP）和路径聚合网络（PAN）。YOLOv4的骨干网络采用CSPDarknet53，颈部采用SPP和PAN，头部保持YOLOv3的设计，在COCO数据集上以较低的参数量实现了当时最优的检测性能。')

    add_body_text(doc, 'YOLOv5由Ultralytics团队发布，虽然未发表学术论文，但因其工程化的实现和易用性迅速成为工业界最受欢迎的检测框架之一。YOLOv5提供了n/s/m/l/x多种模型变体，通过深度和宽度缩放因子灵活调节模型规模。该版本引入了自适应锚框计算、自动混合精度训练、模型集成导出等实用功能，降低了非专业人员的使用门槛。')

    add_body_text(doc, 'YOLOv6由美团视觉智能部提出，面向工业检测场景进行了深度优化。该版本设计了EfficientRep骨干网络，将标准卷积替换为重参数化卷积（RepConv），在训练时使用多分支结构提升表达能力，在推理时融合为单分支结构降低延迟。颈部引入了Rep-PAN和CSPStackRep模块，检测头采用Efficient Decoupled Head将分类和回归任务解耦处理。同时提出SimOTA标签分配策略动态匹配正样本。')

    add_body_text(doc, 'YOLOv7由原YOLOv4作者团队提出，在架构设计上进行了多项创新。模型提出了扩展高效层聚合网络（Extended Efficient Layer Aggregation Networks, E-ELAN），在不破坏原有梯度路径的情况下扩展网络容量。引入了基于串联模型的模型缩放方法，同时缩放深度和宽度。提出了辅助训练头（Auxiliary Head）和引导头（Lead Head）的协同训练策略，以及粗到细的标签分配机制。')

    add_body_text(doc, 'YOLOv8由Ultralytics团队在2023年发布，采用解耦检测头将分类和回归分支分离处理，提升了收敛速度。骨干网络采用C2f模块替代了CSP结构中的C3模块，通过增加跨层连接数量增强梯度和信息流动。采用无锚框检测方式简化了后处理流程，同时引入了任务对齐学习（Task-Aligned Learning）机制。该版本支持检测、分割、姿态估计、分类等多任务，具有良好的生态整合性。')

    add_body_text(doc, 'YOLOv9提出了可编程梯度信息（Programmable Gradient Information, PGI）和广义高效层聚合网络（Generalized Efficient Layer Aggregation Network, GELAN），从信息瓶颈理论的角度分析了深度网络中信息损失的机理。通过辅助可逆分支提供可靠的梯度信息，缓解了深层网络中小目标信息丢失的问题。YOLOv10由清华大学团队提出，首次实现了无需非极大值抑制（NMS-free）的一致双分配策略，通过一对多的训练监督和一对一推理分配的一致性匹配消除了后处理需求。同时引入效率驱动的模型设计原则，在保持高精度的同时显著降低了推理延迟。')

    add_body_text(doc, 'YOLOv11在YOLOv8的基础上进行了系统性优化，在骨干网络中引入了C3k2模块和C2PSA（Cross-Stage Partial with Position-Sensitive Attention）模块。C3k2模块在C2f的基础上提供了可配置的卷积核大小选项，使特征提取更加灵活。C2PSA模块在跨阶段部分连接中嵌入了位置敏感注意力机制，增强了骨干网络对关键空间位置的感知能力。YOLOv11提供了n/s/m/l/x五种模型规模，其中YOLOv11s（945万参数、21.7 GFLOPs）在精度和效率之间取得了很好的平衡，是适合在有限算力条件下进行迁移学习和改进研究的理想基线。')

    # 1.2.3
    add_heading_custom(doc, '1.2.3 注意力机制在目标检测中的应用', level=2)

    add_body_text(doc, '注意力机制源于人类视觉系统的选择性关注特性，其核心思想是引导模型将有限的计算资源聚焦于输入中最具信息量的部分。在深度学习中，注意力机制最早在自然语言处理领域取得突破，随后被广泛引入计算机视觉任务。在目标检测特别是小目标检测中，注意力机制能够有效增强目标区域的特征响应，抑制无关背景信息的干扰。')

    add_body_text(doc, 'Hu等提出的挤压与激励网络（Squeeze-and-Excitation Networks, SENet）通过全局平均池化压缩空间信息，再通过全连接层学习通道间的依赖关系，生成通道维度的注意力权重。SE模块以极小的参数开销（每层增加约2%的参数量）显著提升了网络的特征判别能力，在ImageNet分类竞赛中取得了当年冠军。该机制的简洁高效使其成为最广泛使用的通道注意力方案之一。')

    add_body_text(doc, 'Woo等提出的卷积块注意力模块（Convolutional Block Attention Module, CBAM）在通道注意力的基础上进一步引入了空间注意力分支。通道注意力通过全局平均池化和最大池化的并行的编码方式获取更丰富的全局统计信息，空间注意力沿通道维度进行池化操作生成空间注意力图。CBAM的串行设计（先通道注意力后空间注意力）被证明优于并行设计，在各主流网络架构上均可即插即用，取得了稳定的精度提升。')

    add_body_text(doc, 'Hou等提出的坐标注意力（Coordinate Attention, CA）将位置信息显式编码到通道注意力中。该方法分别沿水平方向和垂直方向进行自适应平均池化，保留一维空间位置信息，然后将两路编码特征进行拼接和卷积处理，生成同时编码通道关系和空间坐标信息的注意力图。CA模块在移动网络上的轻量级设计中表现优异，适合资源受限的小目标检测场景。')

    add_body_text(doc, 'Wang等提出的ECA-Net（Efficient Channel Attention）使用一维卷积替代SE模块中的全连接层进行跨通道交互，避免了降维带来的信息损失。ECA模块的自适应卷积核大小根据通道数自动确定，无需任何人工调参，保持了极低的计算开销。')

    add_body_text(doc, 'Li等提出的EMA（Efficient Multi-Scale Attention）注意力机制通过分组策略和跨空间交互学习实现高效的多尺度注意力。EMA将输入特征沿通道维度均匀分组，在每组内使用两个并行分支：1×1分支沿高度和宽度方向编码空间方向信息，生成门控式注意力权重；3×3分支捕获局部空间上下文。两组特征通过批量矩阵乘法进行交叉空间学习，利用Softmax生成空间注意力分布图。EMA具有参数最少（每个分组内共享参数）、多尺度感知能力强、适合密集小目标场景的优点，是本文采用的注意力机制方案。')

    add_body_text(doc, 'Yang等提出的SimAM（Simple Parameter-Free Attention Module）基于神经科学中的空间抑制理论，利用已知神经元对其邻域神经元产生的能量函数推导注意力权重。SimAM无需任何额外参数即可计算三维注意力权重（通道和空间维度同时加权），具有完全无参数的轻量化优势，但其空间区分能力相对受限。')

    add_body_text(doc, '在目标检测中，注意力机制的插入位置对最终性能有显著影响。现有研究通常将注意力模块放置于骨干网络末端、颈部网络的各层级融合后或检测头之前。不同位置的注意力关注不同层次的特征：骨干网络中的注意力侧重基础特征的增强，颈部融合后的注意力侧重多尺度特征的优化，检测头前的注意力侧重分类和回归前的特征精炼。针对小目标检测的特定需求，在多个层级部署注意力模块以形成递进式的注意力增强链路，是本文的研究重点之一。')

    # 1.2.4
    add_heading_custom(doc, '1.2.4 无人机航拍图像目标检测研究进展', level=2)

    add_body_text(doc, '无人机航拍图像中的目标检测面临目标尺寸微小、密集分布、背景复杂多变、光照条件不均匀等独特挑战，近年来成为国内外计算机视觉领域的研究热点。')

    add_body_text(doc, '在国内研究方面，林世颢等针对无人机航拍图像中目标尺寸小、目标密集、漏检和误检等问题，提出了改进YOLOv11的小目标检测算法GR-YOLOv11n。该算法引入GhostNet v2替换主干网络中的C3k2模块，通过深度可分离卷积和注意力机制提取更为丰富的特征信息；同时设计了C2f-RepNCSPFPN模块优化卷积结构和特征处理流程，显著降低了模型的计算复杂度和参数量。在VisDrone2019数据集上的实验结果表明，GR-YOLOv11n相比YOLOv11n在mAP50上提高了7.6个百分点，mAP50-95上提高了5.4个百分点，在满足轻量化要求的同时有效提升了小目标的检测精度。')

    add_body_text(doc, '贺智轩等针对无人机航拍图像中密集小目标、显著多尺度变化和复杂场景干扰等问题，提出了DMF-YOLOv11算法。该算法设计了双向辅助特征金字塔网络（Dual Bi-directional Auxiliary FPN）作为颈部结构，通过多层次双向特征融合增强极小目标和规则小目标的特征表示；构建了多分支混合卷积模块（Multi-branch Hybrid Convolution），利用并行异构卷积路径提升对小尺度目标的敏感性；将自调制特征聚合网络与主干C3K2模块深度融合，提出C3K2_FMB模块协同提取局部细节与非全局上下文特征。在VisDrone2019数据集上的实验结果表明，DMF-YOLOv11的mAP50和mAP50-95分别达到46.2%和28.4%，较基线YOLOv11n提升了11.5和8.3个百分点，召回率提升9.4个百分点至44.6%。')

    add_body_text(doc, '孙建民等针对无人机目标检测中多尺度特征不足、小目标漏检及算力受限等问题，提出了基于YOLOv11n的高效轻量化改进算法。该研究提出C3k2-S模块，融合C3k2的多尺度特征提取与StarBlock的高维映射能力强化特征建模；采用ADown机制替代传统下采样，通过自适应步幅与池化策略优化特征表达，降低计算量并提升小目标敏感性；提出轻量化HSPAN-C结构，结合多尺度特征融合与空间金字塔注意力增强检测精度；构建LSDECD检测头，利用共享卷积与群归一化减少参数量。实验结果表明，改进后模型的mAP50和mAP50-95分别上升了3.4%和3.1%，参数量和计算量分别降低了46.1%和22.2%，满足复杂场景下无人机的检测需求。')

    add_body_text(doc, '在国际研究方面，X. Wu等提出了基于改进YOLOv11n的轻量化增强解决方案RLD-YOLO。该算法结合RepConv结构重参数化技术，在训练时保留多分支特征表达能力，推理时自动转换为高效的单分支结构；设计LKAConv大核注意力模块，通过7×7深度可分离卷积和空间注意力机制增强小目标特征捕获能力；引入DASI动态自适应融合模块，通过可学习的权重分配优化多尺度特征交互。VisDrone2019-DET数据集上的实验结果表明，融合LKAConv、RepConv和DASI的RLD-YOLO模型，mAP50和mAP50-95分别提升了2.02和1.17个百分点。')

    add_body_text(doc, '综合国内外研究现状可以看出，基于YOLOv11的小目标检测研究呈现以下发展趋势：（1）结构创新驱动性能提升——通过引入注意力机制、优化特征融合结构等不断提升模型表征能力；（2）轻量化与边缘部署并重——面向资源受限设备实现精度与效率的平衡，满足无人机等移动端的实时推理需求；（3）多尺度建模持续改进——P2检测层、跨层密集连接、自适应特征融合等多尺度策略成为提升小目标检测性能的关键手段；（4）多技术融合创新——注意力机制、Transformer、跨层特征交互、重参数化等多元技术路径为小目标检测提供了丰富的解决方案。本文正是在这一研究背景下，综合借鉴上述方法的有益成果，围绕P2检测层、EMA注意力机制和Focal-EIoU损失优化三条技术路线开展系统性改进研究。')

    # ---------- 1.3 ----------
    add_heading_custom(doc, '1.3 主要研究内容', level=1)

    add_body_text(doc, '本文以无人机航拍场景下的小目标检测为研究对象，以VisDrone2019数据集为实验平台，以YOLOv11为基线算法框架，针对现有检测方法在密集小目标场景下特征表达不足、多尺度信息利用不充分、定位精度受限等问题，从模型结构改进、注意力机制引入和损失函数优化三个维度开展系统性研究。主要研究内容包括以下几个方面：')

    add_body_text(doc, '（1）数据集分析与预处理。深入分析VisDrone2019无人机航拍数据集的统计特性，包括类别分布、目标尺寸分布、目标密度分布、锚框先验尺寸等，为模型设计和超参数调优提供数据支撑。完成数据集格式转换和预处理流程，确保与YOLOv11框架的输入规范兼容。')

    add_body_text(doc, '（2）基线模型建立与参数调优。以YOLOv11s为基线模型，在VisDrone2019数据集上进行系统的训练参数调优实验。针对RTX 3060 Laptop 6GB显存的硬件约束，优化批次大小、输入分辨率等内存关键参数。通过对比实验确定最优的优化器（AdamW vs SGD）、学习率策略（余弦退火）、数据增强组合（Mosaic+MixUp+Copy-Paste）和损失权重配置。建立性能稳定的基线模型，为后续改进提供可靠的对比基准。')

    add_body_text(doc, '（3）P2小目标检测层设计。针对YOLOv11标准架构（P3/8、P4/16、P5/32三级检测）在高分辨率小目标检测能力上的不足，设计并添加P2检测层（P2/4、160×160分辨率）。通过在特征金字塔的顶层添加一个上采样模块，将P3层的特征上采样至P2分辨率与骨干网络P2层的浅层特征进行融合，生成高分辨率的P2特征图。同时在自下而上的路径聚合中增加P2到P3的下采样路径，保持双向特征流动。P2检测层的引入使模型获得160×160的高分辨率检测能力，对极小目标的特征感知和定位精度有望显著提升。')

    add_body_text(doc, '（4）EMA注意力机制引入。针对密集小目标场景中背景噪声干扰严重的问题，引入EMA（Efficient Multi-Scale Attention）注意力机制增强关键特征区域的响应。EMA通过分组策略和跨空间交互学习实现高效的多尺度注意力计算：将输入特征沿通道维度均匀分组，在每组内使用1×1分支编码空间方向信息生成门控式权重，使用3×3分支捕获局部空间上下文，两组特征通过批量矩阵乘法进行交叉空间学习生成空间注意力分布图。本文在颈部网络的多个层级（P2、P3、P4）部署EMA模块，形成递进式注意力增强链路。对比CBAM和CA等注意力方案，分析EMA在小目标检测任务上的性能优势。')

    add_body_text(doc, '（5）Focal-EIoU损失函数优化。针对小目标检测中正负样本不平衡和定位精度要求高的问题，研究改进边界框回归损失函数。引入Focal-EIoU损失机制，在标准EIoU（Efficient IoU）损失的基础上添加聚焦因子，动态降低已良好回归锚框的损失贡献，使训练过程更关注定位精度差的小目标样本。同时，通过系统的消融实验验证各改进模块的独立贡献和协同效应。')

    add_body_text(doc, '（6）对比实验与性能评估。在VisDrone2019数据集上与YOLOv8s、YOLOv10s、RT-DETR等主流检测算法进行公平对比，从mAP50、mAP50-95、精确率、召回率、参数量、计算量（GFLOPs）、推理速度（FPS）等多个维度评估改进模型的综合性能。通过消融实验分别验证P2检测层、EMA注意力机制和Focal-EIoU损失函数的独立贡献。')

    add_body_text(doc, '（7）可视化检测系统设计与实现。基于PyQt5框架设计并实现一个小目标可视化检测系统，将训练好的改进模型部署为可交互的桌面应用程序。系统支持图片文件、视频文件、摄像头实时采集等多种输入方式，提供置信度阈值、IoU阈值等检测参数的实时调节功能，并将检测结果（类别标签、置信度、边界框）实时叠加显示。通过系统设计与开发，将算法研究成果转化为可操作的实用工具。')

    # ---------- 1.4 ----------
    add_heading_custom(doc, '1.4 论文结构安排', level=1)

    add_body_text(doc, '本文共分为六章，各章内容安排如下：')

    add_body_text(doc, '第一章：绪论。介绍无人机航拍小目标检测的研究背景与意义，系统综述国内外在小目标检测、YOLO算法演进、注意力机制和无人机航拍目标检测等方面的研究进展，分析当前研究存在的不足，明确本文的研究内容和论文结构。')

    add_body_text(doc, '第二章：相关理论与技术。详细介绍本文研究涉及的基础理论和核心技术，包括卷积神经网络基本组件、YOLOv11模型架构详解（骨干网络、颈部网络、检测头）、注意力机制原理及主流实现方案、小目标检测关键技术和方法、VisDrone2019数据集特征分析，为后续章节的改进研究提供理论支撑。')

    add_body_text(doc, '第三章：改进的YOLOv11小目标检测算法。系统阐述本文提出的改进方案，包括P2小目标检测层的设计与实现、EMA注意力机制的引入策略与部署方案、Focal-EIoU损失函数的优化设计，以及三者的整合架构。详细说明模型配置文件的构建和训练策略的设计。')

    add_body_text(doc, '第四章：实验设计与结果分析。介绍实验环境配置、评价指标体系、训练策略和超参数设置。系统开展消融实验，验证P2检测层、EMA注意力机制和Focal-EIoU损失函数的独立和协同贡献。与YOLOv8s、YOLOv10s、RT-DETR等主流算法进行全面对比，从多个维度评估改进模型的性能优势。对实验结果进行深入分析和讨论。')

    add_body_text(doc, '第五章：小目标检测系统设计与实现。从系统需求分析出发，设计基于PyQt5的可视化检测系统架构，划分功能模块（模型加载模块、数据输入模块、检测推理模块、结果显示模块），分别实现各模块的功能，并进行系统集成和功能测试。')

    add_body_text(doc, '第六章：总结与展望。总结本文的主要研究工作、创新点和取得的成果，分析存在的不足和局限性，展望未来的改进方向和研究计划。')

    # ======================================================================
    # 第二章 相关理论与技术
    # ======================================================================
    doc.add_page_break()
    add_heading_custom(doc, '第二章 相关理论与技术', level=0)

    # ---------- 2.1 ----------
    add_heading_custom(doc, '2.1 卷积神经网络基础', level=1)

    add_body_text(doc, '卷积神经网络（Convolutional Neural Network, CNN）是深度学习在计算机视觉领域取得突破性进展的核心技术架构。CNN通过局部连接、权值共享和下采样三大设计原则，有效解决了传统全连接网络在处理图像等高维数据时面临的参数爆炸和空间信息丢失问题。理解CNN的基本组成和工作原理，是深入掌握现代目标检测算法的基础。')

    add_heading_custom(doc, '2.1.1 卷积层', level=2)

    add_body_text(doc, '卷积层（Convolutional Layer）是CNN的核心组件，通过一组可学习的卷积核（Kernel）在输入特征图上进行滑动窗口计算，提取局部空间特征。卷积操作可以形式化表示为输入特征图X与卷积核W的互相关运算：')

    # Placeholder for formula (will be described in text)
    add_body_text(doc, '其中，(i,j)表示输出特征图的空间坐标，c表示输入通道索引，k和l表示卷积核的空间坐标索引，b表示偏置项。卷积核的尺寸（通常为3×3或1×1）决定了局部感受野的大小，步长（Stride）控制卷积核的滑动间隔，填充（Padding）用于控制输出特征图的空间尺寸。多个不同的卷积核并行工作，各自学习不同类型的特征模式（边缘、纹理、形状等），共同构成卷积层的输出通道。')

    add_body_text(doc, '卷积层的关键优势在于权值共享——同一卷积核在整个输入空间上使用相同的参数，极大地减少了参数量并赋予了网络对平移变化的等变性。在YOLOv11中，标准卷积和深度可分离卷积（Depthwise Separable Convolution）被交替使用，后者将标准卷积分解为逐通道的深度卷积和1×1的逐点卷积两个步骤，有效降低了计算开销。')

    add_heading_custom(doc, '2.1.2 激活函数', level=2)

    add_body_text(doc, '激活函数为神经网络引入非线性变换能力，使网络能够逼近任意复杂的映射关系。现代目标检测网络中最常用的激活函数包括SiLU（Sigmoid Linear Unit，也称Swish）和ReLU（Rectified Linear Unit）。SiLU激活函数的数学形式为SiLU(x)=x·σ(x)，其中σ(x)为Sigmoid函数。SiLU具有平滑的梯度曲线和自门控特性，在深层网络中通常优于ReLU。YOLOv11在骨干网络和颈部网络中广泛采用SiLU作为默认激活函数。')

    add_heading_custom(doc, '2.1.3 批量归一化', level=2)

    add_body_text(doc, '批量归一化（Batch Normalization, BN）通过对每个训练小批量的激活值进行归一化处理，使每层输入的分布保持稳定，有效缓解了深层网络中的内部协变量偏移（Internal Covariate Shift）问题。BN层在每个训练步骤中计算当前小批量的均值和方差，使用可学习的缩放参数γ和偏移参数β对归一化后的值进行线性变换。BN的引入使得网络可以使用更大的学习率、对参数初始化的敏感度降低，并具有一定的正则化效果。在YOLOv11中，每个卷积层之后均紧跟一个BN层和SiLU激活函数，构成Conv-BN-SiLU的标准组合。')

    add_heading_custom(doc, '2.1.4 池化层与下采样', level=2)

    add_body_text(doc, '池化层（Pooling Layer）通过对局部邻域内的特征进行聚合统计，实现空间尺寸的缩减和特征的不变性增强。最大池化（Max Pooling）取局部窗口内的最大值，保留最显著的特征响应；平均池化（Average Pooling）计算局部窗口内的均值，平滑特征表达。空间金字塔池化（Spatial Pyramid Pooling, SPP）使用多个不同尺寸的池化窗口并行处理同一特征图，将各尺度的池化结果沿通道维度拼接，从而融合多尺度感受野信息。YOLOv11骨干网络末端的SPPF（Spatial Pyramid Pooling Fast）模块通过串行连接三个5×5最大池化操作替代并行多尺度池化，以更低的计算开销达到等效的多尺度池化效果。')

    add_heading_custom(doc, '2.1.5 上采样与特征融合', level=2)

    add_body_text(doc, '上采样操作用于将低分辨率特征图恢复至高分辨率，是实现特征金字塔自顶向下路径的关键技术。最近邻插值（Nearest Neighbor Interpolation）直接复制邻近像素值，计算简单且不会引入虚假的高频信息，是YOLO系列特征金字塔中的默认上采样方式。特征融合操作中，通道拼接（Concatenation）将来自不同层级的特征图沿通道维度堆叠，保留各路的完整信息，供后续卷积层学习跨层级的特征组合。YOLOv11颈部网络通过交替进行上采样-拼接-C3k2处理和下采样-拼接-C3k2处理，构建了双向的特征信息流动管道。')

    # ---------- 2.2 ----------
    add_heading_custom(doc, '2.2 YOLOv11模型架构', level=1)

    add_body_text(doc, 'YOLOv11是Ultralytics团队在YOLOv8基础上迭代优化的检测框架，于2024年发布。该版本在保持YOLO系列简洁高效设计哲学的前提下，在骨干网络、特征融合和训练策略等方面进行了多项重要改进。本节从骨干网络、颈部网络和检测头三个组成部分，详细剖析YOLOv11的网络结构。')

    add_heading_custom(doc, '2.2.1 骨干网络', level=2)

    add_body_text(doc, 'YOLOv11的骨干网络负责从输入图像中逐层提取多尺度的语义特征。输入图像（640×640×3）首先经过两个步长为2的3×3标准卷积层，空间分辨率依次降至320×320（64通道）和160×160（128通道）。随后交替使用C3k2模块和标准卷积下采样操作构建层级化的特征提取流水线。')

    add_body_text(doc, 'C3k2模块是YOLOv11骨干网络的核心构建单元，是在C2f模块基础上的重要改进。C2f（Cross-Stage Partial with Two Convolutions）的设计思想源于CSPNet的跨阶段部分连接策略：输入特征首先通过一个1×1卷积（cv1）进行通道压缩和变换，然后将输出沿通道维度均分为两路，一路直接传递至输出端（保留原始信息），另一路通过多个级联的Bottleneck模块进行深度特征提取，最后将两路特征沿通道维度拼接并通过第二个1×1卷积（cv2）进行融合。C3k2在C2f的基础上引入了可配置的卷积核大小选项（c3k参数）：当c3k=True时，Bottleneck中的标准3×3卷积被替换为可配置核大小的C3k变体，增强了对不同尺度特征的适应能力。在大型变体（YOLOv11m/l/x）中，c3k默认启用。')

    add_body_text(doc, '骨干网络的层级结构如下：第0层为P1/2特征（320×320, 32通道），第2层为P2/4特征（160×160, 128通道），第4层为P3/8特征（80×80, 256通道），第6层为P4/16特征（40×40, 256通道），第8-10层为P5/32特征（20×20, 512通道）。在P5层级末端，SPPF模块通过串行最大池化聚合多尺度感受野，C2PSA模块通过位置敏感注意力增强骨干网络的终端特征质量。')

    add_heading_custom(doc, '2.2.2 颈部网络', level=2)

    add_body_text(doc, 'YOLOv11的颈部网络采用改进的PANet结构，在特征金字塔网络（FPN）的基础上增加了自下而上的路径聚合分支。FPN自顶向下路径通过逐级上采样和跨层拼接，将深层强语义特征传播至浅层高分辨率层；PAN自底向上路径通过逐级下采样和拼接，将浅层精确位置信息传递至深层。双向特征流动使每个检测层级都融合了来自多个尺度的丰富信息。')

    add_body_text(doc, '标准YOLOv11的颈部网络生成三个检测层级的特征图：P3/8层（80×80, 128通道，负责检测小目标）、P4/16层（40×40, 256通道，负责检测中等目标）、P5/32层（20×20, 512通道，负责检测大目标）。每个融合节点由一个C3k2模块完成跨层信息的整合和精炼。各检测层级之间的连接索引经过精心设计，确保信息流动的连续性。')

    add_heading_custom(doc, '2.2.3 检测头', level=2)

    add_body_text(doc, 'YOLOv11采用无锚框（Anchor-Free）的检测范式，检测头在每个特征图位置直接预测目标的边界框属性和类别概率，省去了复杂的锚框匹配和后处理步骤。检测头由两个并行的卷积分支组成：边界框回归分支（cv2）和类别分类分支（cv3）。')

    add_body_text(doc, '边界框回归分支接收特征图（如P3层级128通道），经过两层3×3标准卷积将通道数降至合适的中间维度（c2=max(16, c1/4, reg_max×4)），然后通过一个1×1卷积输出4×reg_max个通道，分别编码边界框的左上角和右下角相对于网格中心点的偏移量分布。在训练阶段，使用Distribution Focal Loss（DFL）对偏移量的离散分布进行监督；在推理阶段，取偏移量分布的加权平均值作为最终回归结果。')

    add_body_text(doc, '类别分类分支采用深度可分离卷积降低参数量：特征图（如128通道）首先通过3×3深度可分离卷积（组数等于通道数）进行空间特征提取，再通过1×1卷积进行通道降维，接着经过第二层3×3深度可分离卷积和1×1卷积，最终输出nc（类别数，VisDrone为10）个通道的类别预测分数。这种设计在保持分类精度的同时显著减少了检测头的参数量和计算量。')

    add_heading_custom(doc, '2.2.4 标签分配与损失函数', level=2)

    add_body_text(doc, 'YOLOv11采用任务对齐学习策略（Task-Aligned Assigner）进行正负样本的标签分配。该策略综合考虑分类得分和定位精度的对齐程度，选择对齐程度最高的top_k个预测作为正样本。任务对齐度量定义为分类得分与IoU的乘积的幂次加权：t=s^α×u^β，其中s为分类得分，u为IoU值，α和β为平衡两个任务的超参数。')

    add_body_text(doc, '损失函数由三部分组成：边界框回归损失（采用CIoU或EIoU Loss）、分布聚焦损失（Distribution Focal Loss, DFL，用于监督边界框偏移量的离散概率分布）和类别分类损失（采用二元交叉熵损失，Binary Cross-Entropy, BCE）。总损失为三者的加权和：Loss=λ_box×L_box+λ_dfl×L_dfl+λ_cls×L_cls。其中λ_box、λ_dfl和λ_cls为对应的损失权重系数，用于调节各任务在联合优化中的相对重要性。')

    # ---------- 2.3 ----------
    add_heading_custom(doc, '2.3 注意力机制原理', level=1)

    add_body_text(doc, '注意力机制是深度学习中最具影响力的技术创新之一，其核心思想源自人类认知过程中的选择性注意。在面对复杂的视觉场景时，人类视觉系统并非平等对待所有区域，而是自动聚焦于最具信息量的部分，忽略无关的背景信息。深度注意力机制模拟了这一认知能力，通过可学习的权重分布对输入特征的不同部分进行差异化加权，实现计算资源的自适应分配。')

    add_heading_custom(doc, '2.3.1 通道注意力', level=2)

    add_body_text(doc, '通道注意力机制聚焦于特征图通道维度的重要性建模，通过全局信息聚合学习不同通道的相对重要性权重。SE模块是通道注意力的开创性工作，其计算流程为：首先通过全局平均池化将C×H×W的特征图压缩为C×1×1的通道描述符；然后通过两个全连接层（附带ReLU和Sigmoid激活）将通道描述符映射为通道注意力权重；最后将权重与原始特征图对应通道相乘，实现特征的通道级重标定。SE模块的设计蕴含了"使用全局信息选择性地强调有用特征、抑制无用特征"的深刻思想。')

    add_body_text(doc, 'ECA模块进一步简化了通道注意力的计算：使用一维卷积替代全连接层实现局部跨通道交互，卷积核的大小k根据通道数C自适应确定（k=ψ(C)，其中ψ(C)=|log2(C)/γ+b/γ|_odd），无需手动降维调参。实验表明，适当的跨通道交互覆盖范围（核大小k）对注意力效果至关重要，过大的交互范围可能引入冗余，而过小则无法建立有效的通道依赖关系。')

    add_heading_custom(doc, '2.3.2 空间注意力', level=2)

    add_body_text(doc, '空间注意力机制关注特征图空间维度的重要性分布，为不同空间位置分配差异化的权重。与通道注意力关注"哪些通道更重要"不同，空间注意力关注"哪些位置更重要"。空间注意力通常沿通道维度进行压缩（如最大池化和平均池化），得到1×H×W的空间描述符，然后通过卷积操作生成空间注意力图。对于小目标检测而言，空间注意力能够引导网络聚焦于目标可能出现的少数关键区域，有效抑制大面积无关背景的干扰。')

    add_heading_custom(doc, '2.3.3 EMA注意力机制', level=2)

    add_body_text(doc, 'EMA（Efficient Multi-Scale Attention）是本文采用的核心注意力机制。EMA通过分组策略和跨空间交互学习，在极低的参数开销下实现高效的多尺度特征增强。EMA的完整计算流程包含以下步骤：')

    add_body_text(doc, '（1）通道分组：将输入特征张量X∈R^(B×C×H×W)沿通道维度均匀划分为g组，每组包含cg=C/g个通道。将分组维度展开为批次维度，得到Xg∈R^((B·g)×cg×H×W)。分组操作使不同通道组的注意力计算相互独立，各组学习差异化的空间关注模式。')

    add_body_text(doc, '（2）并行多尺度分支：每组特征分别通过两个并行的卷积分支进行处理。1×1分支首先沿高度和宽度方向分别进行平均池化，得到沿H方向的一维特征（（B·g）×cg×H×1）和沿W方向的一维特征（（B·g）×cg×1×W）。将两者拼接为（（B·g）×cg×（H+W）×1 ）的联合张量，通过1×1卷积和SiLU激活进行跨通道交互，再拆分回H和W两个方向，分别通过1×1卷积和Sigmoid生成门控权重。原始分组特征乘以两个方向的门控权重，再通过GroupNorm归一化。3×3分支直接通过3×3标准卷积和SiLU激活提取局部空间上下文。')

    add_body_text(doc, '（3）跨空间学习：将1×1分支和3×3分支的输出分别沿空间维度展平，通过Softmax函数生成空间注意力分布矩阵A1和A3∈R^((B·g)×cg×(HW))。同时计算两个分支各自的全局平均池化向量v1、v3∈R^((B·g)×1×cg)。通过批量矩阵乘法，v1与A3相乘得到空间注意力图s1（利用3×3分支的空间分布信息调制1×1分支的全局语义），v3与A1相乘得到s2（利用1×1分支的空间分布信息调制3×3分支的局部语义）。最终的空间注意力图为s=σ(s1+s2)。')

    add_body_text(doc, '（4）特征重标定：将空间注意力图s与原始分组特征逐元素相乘，完成空间维度的特征重标定。最后将分组维度恢复为通道维度，输出增强后的特征。')

    add_body_text(doc, 'EMA注意力机制具有以下突出优势：使用分组而非全通道计算实现了高效的特征并行处理；1×1和3×3双分支设计同时捕获了全局空间依赖和局部细节信息；跨空间学习通过批量矩阵乘法建立了两个分支之间的信息交互通道；参数高度共享使模块的总参数量极其有限。这些特性使EMA特别适合小目标密集场景，能够以极小的额外开销为网络提供高质量的注意力引导。')

    # ---------- 2.4 ----------
    add_heading_custom(doc, '2.4 小目标检测关键技术', level=1)

    add_body_text(doc, '小目标检测之所以困难，根源在于目标尺寸极小导致的可辨识信息不足。本节从特征表达、尺度感知和损失优化三个角度，系统梳理当前主流的小目标检测关键技术。')

    add_heading_custom(doc, '2.4.1 多尺度特征增强', level=2)

    add_body_text(doc, '特征金字塔网络（FPN）是多尺度特征增强的基石性工作。FPN通过自顶向下的上采样路径和横向连接，将深层语义信息传播至浅层高分辨率层，使每个检测层级都蕴含丰富的多尺度上下文。然而，标准FPN仅使用P3-P7共5个层级（对应8-128倍下采样），对于无人机航拍图像中像素面积小于8×8的极小目标而言，P3层级的80×80分辨率仍然不足。P2检测层的引入（对应4倍下采样，160×160分辨率）为极小目标检测提供了更高分辨率的特征平台，在VisDrone等小目标密集数据集上被多项研究证实有效。')

    add_body_text(doc, '除扩展检测层级外，加权特征融合（BiFPN中的快速归一化融合）、自适应特征融合（ASFF中的空间滤波权重）、密集跨层连接（Dense-FPN中的全连接信息传递）等策略从不同角度改进了特征融合的质量。本文在标准PANet的基础上添加P2检测层，同时在P2/P3/P4的融合节点嵌入EMA注意力模块，是对多尺度特征增强方向的进一步探索。')

    add_heading_custom(doc, '2.4.2 损失函数设计', level=2)

    add_body_text(doc, 'IoU系列损失函数是目标检测边界框回归的主流选择。IoU损失直接优化预测框与真实框的交并比，具有尺度不敏感的特征，但其在预测框与真实框无重叠时存在梯度为零的问题。GIoU（Generalized IoU）通过引入包围两框的最小闭包框来提供非重叠情况下的梯度信息。DIoU（Distance IoU）在此基础上增加了中心点距离惩罚项，加速了预测框向目标框的收敛。CIoU（Complete IoU）进一步将长宽比的一致性纳入损失计算，通过引入长宽比惩罚项v=4/π²·(arctan(w_gt/h_gt)-arctan(w/h))²来确保预测框的形状与目标框一致，但由于v未考虑长宽比定义的模糊性（w和h的互换不影响arctan比值），在某些情况下可能导致梯度不稳定。EIoU（Efficient IoU）将长宽比惩罚分解为宽度和高度的独立惩罚项，通过直接回归w和h的差异解决了CIoU长宽比定义模糊的问题，在多个基准上取得了更优的收敛效果。')

    add_body_text(doc, 'Focal Loss的聚焦思想通过调制因子(1-p_t)^γ降低易分类样本的损失贡献，使训练过程聚焦于难分类样本。将这一思想迁移至边界框回归：引入聚焦因子(1-IoU)^γ动态降低高IoU（已良好回归）框的损失权重，加大低IoU（回归质量差）框的损失贡献。Focal-EIoU正是将Focal Loss的聚焦机制与EIoU的精准框回归结合，通过调节γ参数（本文取γ=0.5）平衡聚焦强度与训练稳定性。')

    add_heading_custom(doc, '2.4.3 数据增强策略', level=2)

    add_body_text(doc, '数据增强是提升小目标检测性能的有效手段。Mosaic数据增强将四张训练图像随机裁剪并拼接为一张合成图像，不仅增加了小目标的出现频率，还在拼接边界处引入了更多样的上下文组合。MixUp通过两张图像的线性混合αx1+(1-α)x2生成虚拟训练样本，起到正则化效果。复制粘贴增强（Copy-Paste）将目标实例从源图像剪裁出来并粘贴到目标图像的合适位置，直接增加小目标样本的数量和上下文覆盖。在训练后期（通常最后10-15个epoch）关闭马赛克增强，使模型在真实数据分布上进行精细微调，是当前目标检测训练的通行做法。')

    # ---------- 2.5 ----------
    add_heading_custom(doc, '2.5 VisDrone2019数据集', level=1)

    add_body_text(doc, 'VisDrone2019数据集是由天津大学机器学习与数据挖掘实验室在ECCV 2018研讨会中发起的大规模无人机视觉挑战赛的基准数据集。该数据集使用无人机平台在中国14个不同城市采集，覆盖了城市、郊区、乡村等多种地理环境和晴天、阴天、夜间等多种光照条件。数据集原始图像分辨率从960×540到2000×1500不等，反映了真实无人机航拍场景下的图像质量差异。')

    add_body_text(doc, 'VisDrone2019-DET（目标检测子集）包含6471张训练图像、548张验证图像和1610张测试图像。数据标注涵盖10个常见交通参与类别：行人（pedestrian）、人（people）、自行车（bicycle）、汽车（car）、货车（van）、卡车（truck）、三轮车（tricycle）、带篷三轮车（awning-tricycle）、公交车（bus）和摩托车（motor）。数据集共标注了超过260万个目标边界框，平均每张训练图像包含约400个目标实例，目标密度远高于COCO和PASCAL VOC等通用数据集。')

    add_body_text(doc, '根据统计，VisDrone2019中超过80%的目标像素面积小于32×32（COCO小目标定义），超过50%的目标像素面积小于16×16。汽车（car）和行人（pedestrian）是样本数量最多的两个类别，分别占总标注量的约35%和22%，而带篷三轮车（awning-tricycle）等小众类别的样本量相对稀少，存在显著的类别不平衡问题。图像中存在大量密集遮挡场景，例如拥堵路口的密集车辆群、人行横道上的密集人群，同一区域内目标的IoU重叠度常常超过0.5。这些数据特性使VisDrone2019成为评估小目标检测算法在真实复杂场景下性能的理想测试平台。')

    add_body_text(doc, '在数据预处理方面，VisDrone2019的原始标注格式为(bbox_left, bbox_top, bbox_width, bbox_height, score, object_category, truncation, occlusion)，需要转换为YOLO格式(class_id, x_center, y_center, width, height)（归一化坐标）。Ultralytics框架提供了自动的数据格式转换脚本，在首次加载数据集时自动完成转换并缓存标注文件，简化了数据准备流程。')

    print("=" * 60)
    print("Chapter 1-2 generated successfully.")

    # Save document
    output_path = "E:/yolov11/毕业论文_正文.docx"
    doc.save(output_path)
    print(f"Thesis document saved to: {output_path}")
    print("=" * 60)


if __name__ == '__main__':
    create_thesis()
